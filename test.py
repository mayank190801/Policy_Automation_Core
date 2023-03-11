#for finding the table after the required paragraph and deleting and adding a new table

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

def get_table_below_given_tag(doc, tag):
    """
    This function simply returns us the reference to the table below the tags given to us. 
    First find the paragraph, then find the very next table corresponding to the paragraph
    """
    found = False
    for block in iter_block_items(document):
        # print(block.text if isinstance(block, Paragraph) else '<table>')
        if(isinstance(block, Paragraph)):
            if tag in block.text:
                found = True
        else:
            if(found == True):
                return block            
    return None
                
def display_table_data(table):
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            row_data.append(cell.text)
        print('\t'.join(row_data))
    print('')

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

# template = Document("template.docx")
# for style in template.styles: 
#     print(style)
# print(template.styles)
# custom_table_style = template.styles[_TableStyle('Table Grid')]
document = Document('file.docx')
table = get_table_below_given_tag(document, "search word")
# table_style = table.style
# print(table.style.name)


#Removing the table 
table._element.getparent().remove(table._element)
table._element = None

#experimenting the complete procedure here 
paragraph = document.paragraphs[0];  # however you get this paragraph
table = document.add_table(rows=1, cols=3 )

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# table.style = '_Style 10'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

#Adding a new table below a paragraph and removing that paragraph   
move_table_after(table, paragraph)
paragraph.text = ""

document.save("test.docx")
