from docx import Document
document = Document("file.docx")

def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)

paragraph = document.paragraphs[0];  # however you get this paragraph
table = document.add_table(rows=1, cols=3)
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# table.style = 'Table Grid'
table.style = 'LightShading-Accent1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
   
move_table_after(table, paragraph)
paragraph.text = ""

document.save('test1.docx')      



